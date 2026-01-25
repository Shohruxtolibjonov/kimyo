import logging
from telegram import Update, ReplyKeyboardMarkup, ReplyKeyboardRemove, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes, ConversationHandler, CallbackQueryHandler
from docx import Document
from docx.shared import Pt, RGBColor
from io import BytesIO
import sqlite3
from datetime import datetime

# Logging sozlamalari
logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)

# Bot tokeni va Admin ID
BOT_TOKEN = "8349883706:AAHndp5Ps5NfeBniH0XLIskzbYXTLvEOt5M"
ADMIN_ID = 1365319493

# Conversation holatlar - Foydalanuvchilar uchun
(FISH, MANZIL, SINF, AVVAL_OQIGAN, OTA_ONA, 
 TELEFON, OTA_ONA_TELEFON, KELGAN_SANA, MAQSAD) = range(9)

# Admin holatlar
ADMIN_BROADCAST = 100

# Database sozlash
def init_db():
    """Database yaratish"""
    conn = sqlite3.connect('users.db')
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS users
                 (user_id INTEGER PRIMARY KEY,
                  username TEXT,
                  fish TEXT,
                  join_date TEXT)''')
    conn.commit()
    conn.close()

def add_user(user_id, username, fish):
    """Yangi foydalanuvchini bazaga qo'shish"""
    conn = sqlite3.connect('users.db')
    c = conn.cursor()
    try:
        c.execute("INSERT OR REPLACE INTO users (user_id, username, fish, join_date) VALUES (?, ?, ?, ?)",
                  (user_id, username, fish, datetime.now().strftime('%Y-%m-%d %H:%M:%S')))
        conn.commit()
    except Exception as e:
        logger.error(f"Bazaga qo'shishda xatolik: {e}")
    finally:
        conn.close()

def get_all_users():
    """Barcha foydalanuvchilarni olish"""
    conn = sqlite3.connect('users.db')
    c = conn.cursor()
    c.execute("SELECT user_id, username, fish FROM users")
    users = c.fetchall()
    conn.close()
    return users

def get_users_count():
    """Foydalanuvchilar sonini olish"""
    conn = sqlite3.connect('users.db')
    c = conn.cursor()
    c.execute("SELECT COUNT(*) FROM users")
    count = c.fetchone()[0]
    conn.close()
    return count

async def admin_panel(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Admin panel"""
    user_id = update.effective_user.id
    
    if user_id != ADMIN_ID:
        await update.message.reply_text("❌ Sizda admin huquqi yo'q!")
        return
    
    users_count = get_users_count()
    
    keyboard = [
        [InlineKeyboardButton("📊 Statistika", callback_data='admin_stats')],
        [InlineKeyboardButton("📢 Xabar yuborish", callback_data='admin_broadcast')],
        [InlineKeyboardButton("👥 Foydalanuvchilar ro'yxati", callback_data='admin_users_list')],
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await update.message.reply_text(
        f"🔐 <b>ADMIN PANEL</b>\n\n"
        f"👥 Jami foydalanuvchilar: <b>{users_count}</b>\n"
        f"📅 Bugun: {datetime.now().strftime('%d.%m.%Y')}\n\n"
        f"Quyidagi tugmalardan birini tanlang:",
        reply_markup=reply_markup,
        parse_mode='HTML'
    )

async def admin_callback(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Admin tugmalarini boshqarish"""
    query = update.callback_query
    await query.answer()
    
    if query.data == 'admin_stats':
        users_count = get_users_count()
        await query.edit_message_text(
            f"📊 <b>STATISTIKA</b>\n\n"
            f"👥 Jami foydalanuvchilar: <b>{users_count}</b>\n"
            f"📅 Bugun: {datetime.now().strftime('%d.%m.%Y')}",
            parse_mode='HTML'
        )
        
    elif query.data == 'admin_broadcast':
        await query.message.reply_text(
            "📢 <b>XABAR YUBORISH</b>\n\n"
            "Barcha foydalanuvchilarga yubormoqchi bo'lgan xabaringizni yozing:\n\n"
            "Bekor qilish uchun: /cancel",
            parse_mode='HTML'
        )
        return ADMIN_BROADCAST
        
    elif query.data == 'admin_users_list':
        users = get_all_users()
        if not users:
            await query.edit_message_text("❌ Hali foydalanuvchilar yo'q!")
            return
        
        users_text = "👥 <b>FOYDALANUVCHILAR RO'YXATI</b>\n\n"
        for i, (user_id, username, fish) in enumerate(users, 1):
            users_text += f"{i}. {fish or 'Nomsiz'} (@{username or 'username_yoq'})\n"
        
        # Uzun bo'lsa, fayl sifatida yuborish
        if len(users_text) > 4000:
            file = BytesIO(users_text.encode('utf-8'))
            file.name = 'users_list.txt'
            await query.message.reply_document(
                document=file,
                filename='foydalanuvchilar.txt',
                caption="📋 Barcha foydalanuvchilar ro'yxati"
            )
        else:
            await query.edit_message_text(users_text, parse_mode='HTML')
    
    return ConversationHandler.END

async def broadcast_message(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Barcha foydalanuvchilarga xabar yuborish"""
    message_text = update.message.text
    users = get_all_users()
    
    success_count = 0
    failed_count = 0
    
    status_msg = await update.message.reply_text("📤 Xabar yuborilmoqda...")
    
    for user_id, username, fish in users:
        try:
            await context.bot.send_message(
                chat_id=user_id,
                text=f"📢 <b>ADMIN XABARI</b>\n\n{message_text}",
                parse_mode='HTML'
            )
            success_count += 1
        except Exception as e:
            logger.error(f"Foydalanuvchi {user_id} ga xabar yuborishda xatolik: {e}")
            failed_count += 1
    
    await status_msg.edit_text(
        f"✅ <b>Xabar yuborildi!</b>\n\n"
        f"✅ Muvaffaqiyatli: {success_count}\n"
        f"❌ Xatolik: {failed_count}",
        parse_mode='HTML'
    )
    
    return ConversationHandler.END

async def cancel_broadcast(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Xabar yuborishni bekor qilish"""
    await update.message.reply_text("❌ Xabar yuborish bekor qilindi.")
    return ConversationHandler.END

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Botni boshlash va foydalanuvchini kutib olish"""
    user = update.effective_user
    
    # Adminni tekshirish
    if user.id == ADMIN_ID:
        keyboard = [
            [InlineKeyboardButton("🔐 Admin Panel", callback_data='show_admin_panel')],
            [InlineKeyboardButton("📝 Ariza to'ldirish", callback_data='start_form')]
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        
        await update.message.reply_text(
            f"Assalomu alaykum, Admin!\n\n"
            "Quyidagi tugmalardan birini tanlang:",
            reply_markup=reply_markup
        )
        return ConversationHandler.END
    
    await update.message.reply_text(
        f"Assalomu alaykum, {user.first_name}!\n\n"
        "🧪 Kimyo kursimizga yozilishingiz uchun so'rovnomalarga "
        "aniq va to'la javob berishingiz kerak.\n\n"
        "📝 Iltimos, to'liq ismingizni kiriting (F.I.SH):",
        reply_markup=ReplyKeyboardRemove()
    )
    
    # Foydalanuvchi ma'lumotlarini saqlash uchun
    context.user_data['user_info'] = {
        'telegram_username': user.username or 'Username yo\'q',
        'telegram_id': user.id
    }
    
    return FISH

async def start_form_callback(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Admin uchun ariza to'ldirish boshlash"""
    query = update.callback_query
    await query.answer()
    
    user = query.from_user
    
    await query.message.reply_text(
        f"Assalomu alaykum, {user.first_name}!\n\n"
        "🧪 Kimyo kursimizga yozilishingiz uchun so'rovnomalarga "
        "aniq va to'la javob berishingiz kerak.\n\n"
        "📝 Iltimos, to'liq ismingizni kiriting (F.I.SH):",
        reply_markup=ReplyKeyboardRemove()
    )
    
    context.user_data['user_info'] = {
        'telegram_username': user.username or 'Username yo\'q',
        'telegram_id': user.id
    }
    
    return FISH

async def show_admin_panel_callback(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Callback orqali admin panelni ko'rsatish"""
    query = update.callback_query
    await query.answer()
    
    users_count = get_users_count()
    
    keyboard = [
        [InlineKeyboardButton("📊 Statistika", callback_data='admin_stats')],
        [InlineKeyboardButton("📢 Xabar yuborish", callback_data='admin_broadcast')],
        [InlineKeyboardButton("👥 Foydalanuvchilar ro'yxati", callback_data='admin_users_list')],
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await query.message.reply_text(
        f"🔐 <b>ADMIN PANEL</b>\n\n"
        f"👥 Jami foydalanuvchilar: <b>{users_count}</b>\n"
        f"📅 Bugun: {datetime.now().strftime('%d.%m.%Y')}\n\n"
        f"Quyidagi tugmalardan birini tanlang:",
        reply_markup=reply_markup,
        parse_mode='HTML'
    )

async def get_fish(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """F.I.SH ni qabul qilish"""
    context.user_data['user_info']['fish'] = update.message.text
    
    await update.message.reply_text(
        "📍 Qaysi tuman, qaysi qishloqdansiz?"
    )
    return MANZIL

async def get_manzil(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Manzilni qabul qilish"""
    context.user_data['user_info']['manzil'] = update.message.text
    
    await update.message.reply_text(
        "🎓 Nechanchi sinfsiz yoki bitirganmisiz?"
    )
    return SINF

async def get_sinf(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Sinf ma'lumotini qabul qilish"""
    context.user_data['user_info']['sinf'] = update.message.text
    
    keyboard = [['Ha', 'Yo\'q']]
    reply_markup = ReplyKeyboardMarkup(keyboard, one_time_keyboard=True, resize_keyboard=True)
    
    await update.message.reply_text(
        "📚 Avval kimyo o'qiganmisiz?",
        reply_markup=reply_markup
    )
    return AVVAL_OQIGAN

async def get_avval_oqigan(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Avval o'qigan ma'lumotini qabul qilish"""
    context.user_data['user_info']['avval_oqigan'] = update.message.text
    
    keyboard = [['Ha', 'Yo\'q']]
    reply_markup = ReplyKeyboardMarkup(keyboard, one_time_keyboard=True, resize_keyboard=True)
    
    await update.message.reply_text(
        "👨‍👩‍👦 Ota-onangiz bormi?",
        reply_markup=reply_markup
    )
    return OTA_ONA

async def get_ota_ona(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Ota-ona ma'lumotini qabul qilish"""
    context.user_data['user_info']['ota_ona'] = update.message.text
    
    await update.message.reply_text(
        "📱 O'zingizning telefon raqamingizni kiriting:\n"
        "(Masalan: +998901234567)",
        reply_markup=ReplyKeyboardRemove()
    )
    return TELEFON

async def get_telefon(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Telefon raqamni qabul qilish"""
    context.user_data['user_info']['telefon'] = update.message.text
    
    await update.message.reply_text(
        "📞 Ota-onangizning telefon raqamini kiriting:\n"
        "(Masalan: +998901234567)"
    )
    return OTA_ONA_TELEFON

async def get_ota_ona_telefon(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Ota-ona telefon raqamini qabul qilish"""
    context.user_data['user_info']['ota_ona_telefon'] = update.message.text
    
    await update.message.reply_text(
        "📅 O'quv markaziga qachon kelgansiz?\n"
        "(Masalan: 15-yanvar 2026 yoki Bugun)"
    )
    return KELGAN_SANA

async def get_kelgan_sana(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Kelgan sanani qabul qilish"""
    context.user_data['user_info']['kelgan_sana'] = update.message.text
    
    await update.message.reply_text(
        "🎯 Kimyo o'qishdan maqsadingiz nima?"
    )
    return MAQSAD

async def get_maqsad(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Maqsadni qabul qilish va ma'lumotlarni adminga yuborish"""
    context.user_data['user_info']['maqsad'] = update.message.text
    user_info = context.user_data['user_info']
    
    # Foydalanuvchini bazaga qo'shish
    add_user(
        user_info.get('telegram_id'),
        user_info.get('telegram_username'),
        user_info.get('fish')
    )
    
    # Word hujjat yaratish
    doc = Document()
    
    # Sarlavha
    title = doc.add_heading('KIMYO KURSI - YANGI TALABA MA\'LUMOTLARI', 0)
    title.alignment = 1  # Center
    
    doc.add_paragraph()
    
    # Ma'lumotlarni qo'shish
    data_pairs = [
        ("👤 F.I.SH:", user_info.get('fish', 'N/A')),
        ("📍 Manzil:", user_info.get('manzil', 'N/A')),
        ("🎓 Sinf:", user_info.get('sinf', 'N/A')),
        ("📚 Avval kimyo o'qiganmi:", user_info.get('avval_oqigan', 'N/A')),
        ("👨‍👩‍👦 Ota-onasi bormi:", user_info.get('ota_ona', 'N/A')),
        ("📱 Telefon raqami:", user_info.get('telefon', 'N/A')),
        ("📞 Ota-ona telefon raqami:", user_info.get('ota_ona_telefon', 'N/A')),
        ("📅 Kelgan sana:", user_info.get('kelgan_sana', 'N/A')),
        ("🎯 Maqsadi:", user_info.get('maqsad', 'N/A')),
    ]
    
    for label, value in data_pairs:
        p = doc.add_paragraph()
        p.add_run(label).bold = True
        p.add_run(f" {value}")
    
    doc.add_paragraph()
    doc.add_paragraph("_" * 50)
    
    # Telegram ma'lumotlari
    p = doc.add_paragraph()
    p.add_run("📱 Telegram Username: ").bold = True
    p.add_run(f"@{user_info.get('telegram_username', 'N/A')}")
    
    p = doc.add_paragraph()
    p.add_run("🆔 Telegram ID: ").bold = True
    p.add_run(str(user_info.get('telegram_id', 'N/A')))
    
    # Hujjatni xotiraga saqlash
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    # Faylni adminga yuborish
    file_name = f"Ariza_{user_info.get('fish', 'Nomsiz').replace(' ', '_')}.docx"
    
    try:
        await context.bot.send_document(
            chat_id=ADMIN_ID,
            document=file_stream,
            filename=file_name,
            caption=f"🆕 Yangi ariza keldi!\n\n"
                    f"👤 {user_info.get('fish', 'N/A')}\n"
                    f"📱 @{user_info.get('telegram_username', 'N/A')}\n"
                    f"📅 Kelgan sana: {user_info.get('kelgan_sana', 'N/A')}"
        )
        
        await update.message.reply_text(
            "✅ Arizangiz muvaffaqiyatli qabul qilindi!\n\n"
            "Ma'lumotlaringiz admin ko'rib chiqish uchun yuborildi. "
            "Tez orada siz bilan bog'lanamiz.\n\n"
            "Rahmat! 🙏"
        )
        
    except Exception as e:
        logger.error(f"Xatolik yuz berdi: {e}")
        await update.message.reply_text(
            "❌ Arizani yuborishda xatolik yuz berdi. "
            "Iltimos, qaytadan urinib ko'ring yoki admin bilan bog'laning."
        )
    
    return ConversationHandler.END

async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """So'rovnomani bekor qilish"""
    await update.message.reply_text(
        "❌ So'rovnoma bekor qilindi.\n"
        "Qayta boshlash uchun /start ni bosing.",
        reply_markup=ReplyKeyboardRemove()
    )
    return ConversationHandler.END

def main() -> None:
    """Botni ishga tushirish"""
    # Database yaratish
    init_db()
    
    application = Application.builder().token(BOT_TOKEN).build()
    
    # Admin panel handler
    application.add_handler(CommandHandler('admin', admin_panel))
    
    # Admin callback handler
    application.add_handler(CallbackQueryHandler(show_admin_panel_callback, pattern='show_admin_panel'))
    application.add_handler(CallbackQueryHandler(start_form_callback, pattern='start_form'))
    
    # Broadcast conversation handler
    broadcast_handler = ConversationHandler(
        entry_points=[CallbackQueryHandler(admin_callback, pattern='admin_broadcast')],
        states={
            ADMIN_BROADCAST: [MessageHandler(filters.TEXT & ~filters.COMMAND, broadcast_message)],
        },
        fallbacks=[CommandHandler('cancel', cancel_broadcast)],
    )
    application.add_handler(broadcast_handler)
    
    # Admin statistika va users list callbacks
    application.add_handler(CallbackQueryHandler(admin_callback, pattern='admin_stats|admin_users_list'))
    
    # Conversation handler - Foydalanuvchilar uchun
    conv_handler = ConversationHandler(
        entry_points=[CommandHandler('start', start)],
        states={
            FISH: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_fish)],
            MANZIL: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_manzil)],
            SINF: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_sinf)],
            AVVAL_OQIGAN: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_avval_oqigan)],
            OTA_ONA: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_ota_ona)],
            TELEFON: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_telefon)],
            OTA_ONA_TELEFON: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_ota_ona_telefon)],
            KELGAN_SANA: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_kelgan_sana)],
            MAQSAD: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_maqsad)],
        },
        fallbacks=[CommandHandler('cancel', cancel)],
    )
    
    application.add_handler(conv_handler)
    
    # Botni ishga tushirish
    logger.info("Bot ishga tushirildi...")
    application.run_polling(allowed_updates=Update.ALL_TYPES)

if __name__ == '__main__':
    main()