import logging
from telegram import Update, ReplyKeyboardMarkup, ReplyKeyboardRemove, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes, ConversationHandler, CallbackQueryHandler
from docx import Document
from docx.shared import Pt, RGBColor
from io import BytesIO
import sqlite3
from datetime import datetime
import os

# Logging sozlamalari
logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)

# Bot tokeni va Admin ID
BOT_TOKEN = os.getenv("BOT_TOKEN", "7216956481:AAFTnOaikisEAVT2XQEfAtbmtEVAqzRWsKc")
ADMIN_ID = int(os.getenv("ADMIN_ID", "1365319493"))

# Railway Volume path (persistent storage)
DB_PATH = os.getenv("DB_PATH", "/data/users.db")

# Conversation holatlar - Foydalanuvchilar uchun
(FISH, MANZIL, SINF, AVVAL_OQIGAN, OTA_ONA, 
 TELEFON, OTA_ONA_TELEFON, KELGAN_SANA, RASM, MAQSAD) = range(10)

# Admin holatlar
ADMIN_BROADCAST = 100
ADMIN_MESSAGE = 101

# Global o'zgaruvchi - davomat sessiyasi
attendance_active = False
current_session_id = None

# Database sozlash
def init_db():
    """Database yaratish"""
    try:
        # Agar /data papkasi mavjud bo'lmasa, yaratish
        db_dir = os.path.dirname(DB_PATH)
        if db_dir and not os.path.exists(db_dir):
            os.makedirs(db_dir, exist_ok=True)
            logger.info(f"Database papkasi yaratildi: {db_dir}")
        
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()     
        # Users jadvali
        c.execute('''CREATE TABLE IF NOT EXISTS users
                     (user_id INTEGER PRIMARY KEY,
                      username TEXT,
                      fish TEXT,
                      join_date TEXT)''')
        # Davomat jadvali
        c.execute('''CREATE TABLE IF NOT EXISTS attendance
                     (id INTEGER PRIMARY KEY AUTOINCREMENT,
                      session_id TEXT,
                      user_id INTEGER,
                      username TEXT,
                      fish TEXT,
                      check_in_time TEXT,
                      session_date TEXT)''')
        
        # Murojaatlar jadvali
        c.execute('''CREATE TABLE IF NOT EXISTS messages
                     (id INTEGER PRIMARY KEY AUTOINCREMENT,
                      user_id INTEGER,
                      username TEXT,
                      fish TEXT,
                      message_text TEXT,
                      sent_time TEXT,
                      is_read INTEGER DEFAULT 0)''')
        conn.commit()
        conn.close()
        logger.info(f"Database muvaffaqiyatli yaratildi: {DB_PATH}")
    except Exception as e:
        logger.error(f"Database yaratishda xatolik: {e}", exc_info=True)

def add_user(user_id, username, fish):
    """Yangi foydalanuvchini bazaga qo'shish"""
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    try:
        c.execute("INSERT OR REPLACE INTO users (user_id, username, fish, join_date) VALUES (?, ?, ?, ?)",
                  (user_id, username, fish, datetime.now().strftime('%Y-%m-%d %H:%M:%S')))
        conn.commit()
        logger.info(f"Foydalanuvchi bazaga qo'shildi: {fish} ({user_id})")
    except Exception as e:
        logger.error(f"Bazaga qo'shishda xatolik: {e}")
    finally:
        conn.close()

def get_all_users():
    """Barcha foydalanuvchilarni olish"""
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute("SELECT user_id, username, fish FROM users")
        users = c.fetchall()
        conn.close()
        return users
    except Exception as e:
        logger.error(f"Foydalanuvchilarni olishda xatolik: {e}")
        return []

def get_users_count():
    """Foydalanuvchilar sonini olish"""
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute("SELECT COUNT(*) FROM users")
        count = c.fetchone()[0]
        conn.close()
        return count
    except Exception as e:
        logger.error(f"Foydalanuvchilar sonini olishda xatolik: {e}")
        return 0

def get_user_info(user_id):
    """Foydalanuvchi ma'lumotini olish"""
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute("SELECT username, fish FROM users WHERE user_id = ?", (user_id,))
        result = c.fetchone()
        conn.close()
        return result if result else (None, None)
    except Exception as e:
        logger.error(f"Foydalanuvchi ma'lumotini olishda xatolik: {e}")
        return (None, None)

def save_attendance(session_id, user_id, username, fish):
    """Davomatni saqlash"""
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        
        # Avval tekshiramiz - bu sessiyada oldin kelganmi?
        c.execute("SELECT id FROM attendance WHERE session_id = ? AND user_id = ?", 
                  (session_id, user_id))
        if c.fetchone():
            conn.close()
            return False  # Oldin kelgan
        
        check_in_time = datetime.now().strftime('%H:%M:%S')
        session_date = datetime.now().strftime('%Y-%m-%d')
        
        c.execute("""INSERT INTO attendance 
                     (session_id, user_id, username, fish, check_in_time, session_date) 
                     VALUES (?, ?, ?, ?, ?, ?)""",
                  (session_id, user_id, username, fish, check_in_time, session_date))
        conn.commit()
        conn.close()
        return True
    except Exception as e:
        logger.error(f"Davomatni saqlashda xatolik: {e}")
        return False

def get_attendance_report(session_id):
    """Davomat hisobotini olish"""
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute("""SELECT fish, username, check_in_time 
                     FROM attendance 
                     WHERE session_id = ? 
                     ORDER BY check_in_time""", (session_id,))
        report = c.fetchall()
        conn.close()
        return report
    except Exception as e:
        logger.error(f"Hisobotni olishda xatolik: {e}")
        return []

def save_message(user_id, username, fish, message_text):
    """Murojaatni saqlash"""
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        sent_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        c.execute("""INSERT INTO messages 
                     (user_id, username, fish, message_text, sent_time) 
                     VALUES (?, ?, ?, ?, ?)""",
                  (user_id, username, fish, message_text, sent_time))
        conn.commit()
        conn.close()
        return True
    except Exception as e:
        logger.error(f"Murojaatni saqlashda xatolik: {e}")
        return False

def get_unread_messages():
    """O'qilmagan murojaatlarni olish"""
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute("""SELECT id, user_id, username, fish, message_text, sent_time 
                     FROM messages 
                     WHERE is_read = 0 
                     ORDER BY sent_time DESC""")
        messages = c.fetchall()
        conn.close()
        return messages
    except Exception as e:
        logger.error(f"Murojaatlarni olishda xatolik: {e}")
        return []

def mark_messages_read():
    """Barcha murojaatlarni o'qilgan deb belgilash"""
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute("UPDATE messages SET is_read = 1 WHERE is_read = 0")
        conn.commit()
        conn.close()
    except Exception as e:
        logger.error(f"Murojaatlarni yangilashda xatolik: {e}")

async def admin_panel(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Admin panel"""
    user_id = update.effective_user.id
    
    if user_id != ADMIN_ID:
        await update.message.reply_text("❌ Sizda admin huquqi yo'q!")
        return
    
    users_count = get_users_count()
    unread_count = len(get_unread_messages()) 
    keyboard = [
        [InlineKeyboardButton("📊 Statistika", callback_data='admin_stats')],
        [InlineKeyboardButton("📢 Xabar yuborish", callback_data='admin_broadcast')],
        [InlineKeyboardButton("👥 Foydalanuvchilar", callback_data='admin_users_list')],
        [InlineKeyboardButton("📝 Davomat", callback_data='admin_attendance')],
        [InlineKeyboardButton(f"💬 Murojaatlar ({unread_count})", callback_data='admin_messages')],
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await update.message.reply_text(
        f"🔐 <b>ADMIN PANEL</b>\n\n"
        f"👥 Jami foydalanuvchilar: <b>{users_count}</b>\n"
        f"💬 Yangi murojaatlar: <b>{unread_count}</b>\n"
        f"📅 Bugun: {datetime.now().strftime('%d.%m.%Y')}\n\n"
        f"Quyidagi tugmalardan birini tanlang:",
        reply_markup=reply_markup,
        parse_mode='HTML'
    )

async def admin_callback(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Admin tugmalarini boshqarish"""
    query = update.callback_query
    await query.answer()
    
    if query.data == 'admin_stats':
        users_count = get_users_count()
        unread_count = len(get_unread_messages())
        await query.edit_message_text(
            f"📊 <b>STATISTIKA</b>\n\n"
            f"👥 Jami foydalanuvchilar: <b>{users_count}</b>\n"
            f"💬 Yangi murojaatlar: <b>{unread_count}</b>\n"
            f"📅 Bugun: {datetime.now().strftime('%d.%m.%Y')}",
            parse_mode='HTML'
        )
        
    elif query.data == 'admin_broadcast':
        await query.message.reply_text(
            "📢 <b>XABAR YUBORISH</b>\n\n"
            "Barcha foydalanuvchilarga yubormoqchi bo'lgan xabaringizni yozing.\n"
            "Siz matn, rasm, video yuborishingiz mumkin!\n\n"
            "Bekor qilish uchun: /cancel",
            parse_mode='HTML'
        )
        return ADMIN_BROADCAST
        
    elif query.data == 'admin_users_list':
        users = get_all_users()
        if not users:
            await query.edit_message_text("❌ Hali foydalanuvchilar yo'q!")
            return ConversationHandler.END
        
        users_text = "👥 <b>FOYDALANUVCHILAR RO'YXATI</b>\n\n"
        for i, (user_id, username, fish) in enumerate(users, 1):
            users_text += f"{i}. {fish or 'Nomsiz'} (@{username or 'username_yoq'})\n"
        
        # Uzun bo'lsa, fayl sifatida yuborish
        if len(users_text) > 4000:
            file = BytesIO(users_text.encode('utf-8'))
            file.name = 'users_list.txt'
            await query.message.reply_document(
                document=file,
                filename='foydalanuvchilar.txt',
                caption="📋 Barcha foydalanuvchilar ro'yxati"
            )
        else:
            await query.edit_message_text(users_text, parse_mode='HTML')
    
    elif query.data == 'admin_attendance':
        keyboard = [
            [InlineKeyboardButton("▶️ Davomat boshlash", callback_data='start_attendance')],
            [InlineKeyboardButton("⏹️ Davomat yakunlash", callback_data='end_attendance')],
            [InlineKeyboardButton("◀️ Orqaga", callback_data='back_to_admin')]
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        
        status = "🟢 Faol" if attendance_active else "🔴 Faol emas"
        await query.edit_message_text(
            f"📝 <b>DAVOMAT TIZIMI</b>\n\n"
            f"Holat: {status}\n\n"
            f"Davomat boshlash yoki yakunlash uchun tugmani bosing:",
            reply_markup=reply_markup,
            parse_mode='HTML'
        )
    
    elif query.data == 'admin_messages':
        messages = get_unread_messages()
        if not messages:
            await query.edit_message_text("✅ Yangi murojaatlar yo'q!")
            mark_messages_read()
            return ConversationHandler.END
        
        messages_text = "💬 <b>YANGI MUROJAATLAR</b>\n\n"
        for msg_id, user_id, username, fish, text, sent_time in messages:
            messages_text += f"👤 <b>{fish or 'Nomsiz'}</b> (@{username or 'username_yoq'})\n"
            messages_text += f"📅 {sent_time}\n"
            messages_text += f"💬 {text}\n"
            messages_text += "─" * 30 + "\n\n"
        
        if len(messages_text) > 4000:
            file = BytesIO(messages_text.encode('utf-8'))
            file.name = 'messages.txt'
            await query.message.reply_document(
                document=file,
                filename='murojaatlar.txt',
                caption="💬 Barcha murojaatlar"
            )
        else:
            await query.edit_message_text(messages_text, parse_mode='HTML')
        
        mark_messages_read()
    
    elif query.data == 'start_attendance':
        global attendance_active, current_session_id
        
        if attendance_active:
            await query.answer("⚠️ Davomat allaqachon faol!", show_alert=True)
            return ConversationHandler.END
        
        attendance_active = True
        current_session_id = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        # Barcha foydalanuvchilarga xabar yuborish
        users = get_all_users()
        for user_id, username, fish in users:
            try:
                keyboard = [[InlineKeyboardButton("✅ Men keldim", callback_data='check_in')]]
                reply_markup = InlineKeyboardMarkup(keyboard)
                
                await context.bot.send_message(
                    chat_id=user_id,
                    text="📝 <b>DAVOMAT BOSHLANDI!</b>\n\n"
                         "Iltimos, kelganingizni tasdiqlash uchun tugmani bosing:",
                    reply_markup=reply_markup,
                    parse_mode='HTML'
                )
            except Exception as e:
                logger.error(f"Foydalanuvchi {user_id} ga xabar yuborishda xatolik: {e}")
        
        await query.edit_message_text(
            "✅ Davomat boshlandi!\n\n"
            f"Sessiya ID: {current_session_id}\n"
            "Barcha foydalanuvchilarga xabar yuborildi."
        )
    
    elif query.data == 'end_attendance':
        global attendance_active
        
        if not attendance_active:
            await query.answer("⚠️ Davomat faol emas!", show_alert=True)
            return ConversationHandler.END
        
        attendance_active = False
        
        # Hisobotni tayyorlash
        report = get_attendance_report(current_session_id)
        
        if not report:
            await query.edit_message_text("❌ Hech kim kelmagan!")
            return ConversationHandler.END
        
        report_text = f"📊 <b>DAVOMAT HISOBOTI</b>\n\n"
        report_text += f"📅 Sana: {datetime.now().strftime('%d.%m.%Y')}\n"
        report_text += f"🆔 Sessiya: {current_session_id}\n\n"
        report_text += f"👥 Kelganlar soni: <b>{len(report)}</b>\n\n"
        
        for i, (fish, username, check_time) in enumerate(report, 1):
            report_text += f"{i}. {fish} - ⏰ {check_time}\n"
        
        await query.edit_message_text(report_text, parse_mode='HTML')
    
    elif query.data == 'back_to_admin':
        users_count = get_users_count()
        unread_count = len(get_unread_messages())
        
        keyboard = [
            [InlineKeyboardButton("📊 Statistika", callback_data='admin_stats')],
            [InlineKeyboardButton("📢 Xabar yuborish", callback_data='admin_broadcast')],
            [InlineKeyboardButton("👥 Foydalanuvchilar", callback_data='admin_users_list')],
            [InlineKeyboardButton("📝 Davomat", callback_data='admin_attendance')],
            [InlineKeyboardButton(f"💬 Murojaatlar ({unread_count})", callback_data='admin_messages')],
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        
        await query.edit_message_text(
            f"🔐 <b>ADMIN PANEL</b>\n\n"
            f"👥 Jami foydalanuvchilar: <b>{users_count}</b>\n"
            f"💬 Yangi murojaatlar: <b>{unread_count}</b>\n"
            f"📅 Bugun: {datetime.now().strftime('%d.%m.%Y')}\n\n"
            f"Quyidagi tugmalardan birini tanlang:",
            reply_markup=reply_markup,
            parse_mode='HTML'
        )
    
    return ConversationHandler.END

async def check_in_callback(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Foydalanuvchi 'Men keldim' tugmasini bosganida"""
    query = update.callback_query
    await query.answer()
    
    if not attendance_active:
        await query.edit_message_text("❌ Davomat yakunlangan!")
        return
    
    user_id = query.from_user.id
    username, fish = get_user_info(user_id)
    
    if not fish:
        await query.edit_message_text("❌ Siz ro'yxatdan o'tmagansiz! /start ni bosing.")
        return
    
    success = save_attendance(current_session_id, user_id, username or 'N/A', fish)
    
    if success:
        check_time = datetime.now().strftime('%H:%M:%S')
        await query.edit_message_text(
            f"✅ <b>Davomat qabul qilindi!</b>\n\n"
            f"👤 {fish}\n"
            f"⏰ Kelgan vaqt: {check_time}",
            parse_mode='HTML'
        )
        
        # Adminga xabar berish
        try:
            await context.bot.send_message(
                chat_id=ADMIN_ID,
                text=f"✅ Yangi davomat!\n\n"
                     f"👤 {fish}\n"
                     f"⏰ {check_time}",
                parse_mode='HTML'
            )
        except:
            pass
    else:
        await query.edit_message_text("⚠️ Siz allaqachon davomatda turgan!")
    return ConversationHandler.END

async def broadcast_message(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Barcha foydalanuvchilarga xabar yuborish"""
    users = get_all_users()
    
    if not users:
        await update.message.reply_text("❌ Hali foydalanuvchilar yo'q!")
        return ConversationHandler.END
    
    success_count = 0
    failed_count = 0
    
    status_msg = await update.message.reply_text("📤 Xabar yuborilmoqda...")
    
    # Matn xabari
    if update.message.text:
        message_text = update.message.text
        for user_id, username, fish in users:
            try:
                await context.bot.send_message(
                    chat_id=user_id,
                    text=f"📢 <b>ASLIDDINOVdan xabar</b>\n\n{message_text}",
                    parse_mode='HTML'
                )
                success_count += 1
            except Exception as e:
                logger.error(f"Foydalanuvchi {user_id} ga xabar yuborishda xatolik: {e}")
                failed_count += 1
    # Rasm xabari
    elif update.message.photo:
        photo = update.message.photo[-1]
        caption = update.message.caption or ""
        for user_id, username, fish in users:
            try:
                await context.bot.send_photo(
                    chat_id=user_id,
                    photo=photo.file_id,
                    caption=f"📢 <b>ASLIDDINOVdan xabar</b>\n\n{caption}",
                    parse_mode='HTML'
                )
                success_count += 1
            except Exception as e:
                logger.error(f"Foydalanuvchi {user_id} ga rasm yuborishda xatolik: {e}")
                failed_count += 1
    # Video xabari
    elif update.message.video:
        video = update.message.video
        caption = update.message.caption or ""
        for user_id, username, fish in users:
            try:
                await context.bot.send_video(
                    chat_id=user_id,
                    video=video.file_id,
                    caption=f"📢 <b>ASLIDDINOVdan xabar</b>\n\n{caption}",
                    parse_mode='HTML'
                )
                success_count += 1
            except Exception as e:
                logger.error(f"Foydalanuvchi {user_id} ga video yuborishda xatolik: {e}")
                failed_count += 1
    # Document xabari
    elif update.message.document:
        document = update.message.document
        caption = update.message.caption or ""
        for user_id, username, fish in users:
            try:
                await context.bot.send_document(
                    chat_id=user_id,
                    document=document.file_id,
                    caption=f"📢 <b>ASLIDDINOVdan xabar</b>\n\n{caption}",
                    parse_mode='HTML'
                )
                success_count += 1
            except Exception as e:
                logger.error(f"Foydalanuvchi {user_id} ga fayl yuborishda xatolik: {e}")
                failed_count += 1
    
    await status_msg.edit_text(
        f"✅ <b>Xabar yuborildi!</b>\n\n"
        f"✅ Muvaffaqiyatli: {success_count}\n"
        f"❌ Xatolik: {failed_count}",
        parse_mode='HTML'
    )
    
    return ConversationHandler.END

async def contact_admin_start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Adminga murojaat boshlash"""
    query = update.callback_query
    await query.answer()
    
    await query.message.reply_text(
        "💬 <b>ADMINGA MUROJAAT</b>\n\n"
        "Xabaringizni yozing. Admin tez orada javob beradi.\n\n"
        "Bekor qilish uchun: /cancel",
        parse_mode='HTML'
    )
    
    return ADMIN_MESSAGE

async def receive_admin_message(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Foydalanuvchidan xabar qabul qilish"""
    user_id = update.effective_user.id
    username, fish = get_user_info(user_id)
    
    message_text = update.message.text
    
    if save_message(user_id, username or 'N/A', fish or 'Nomsiz', message_text):
        await update.message.reply_text(
            "✅ Xabaringiz adminga yuborildi!\n"
            "Tez orada javob beriladi.",
            reply_markup=ReplyKeyboardRemove()
        )
        
        # Adminga xabar yuborish
        try:
            await context.bot.send_message(
                chat_id=ADMIN_ID,
                text=f"💬 <b>YANGI MUROJAAT!</b>\n\n"
                     f"👤 {fish or 'Nomsiz'}\n"
                     f"📱 @{username or 'username_yoq'}\n"
                     f"📅 {datetime.now().strftime('%d.%m.%Y %H:%M')}\n\n"
                     f"💬 {message_text}",
                parse_mode='HTML'
            )
        except Exception as e:
            logger.error(f"Adminga xabar yuborishda xatolik: {e}")
    else:
        await update.message.reply_text("❌ Xatolik yuz berdi. Qayta urinib ko'ring.")
    
    return ConversationHandler.END

async def cancel_broadcast(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Xabar yuborishni bekor qilish"""
    await update.message.reply_text("❌ Xabar yuborish bekor qilindi.")
    return ConversationHandler.END

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Botni boshlash"""
    user = update.effective_user
    
    # Adminni tekshirish
    if user.id == ADMIN_ID:
        keyboard = [
            [InlineKeyboardButton("🔐 Admin Panel", callback_data='show_admin_panel')],
            [InlineKeyboardButton("📝 Ariza to'ldirish", callback_data='start_form')]
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        
        await update.message.reply_text(
            f"Assalomu alaykum, Admin!\n\n"
            "Quyidagi tugmalardan birini tanlang:",
            reply_markup=reply_markup
        )
        return ConversationHandler.END
    
    keyboard = [
        [InlineKeyboardButton("💬 Adminga murojaat", callback_data='contact_admin')]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text(
        f"Assalomu alaykum, {user.first_name}!\n\n"
        "🧪 Kimyo kursimizga yozilishingiz uchun so'rovnomalarga "
        "aniq va to'la javob berishingiz kerak.\n\n"
        "📝 Iltimos, to'liq ismingizni kiriting (F.I.SH):",
        reply_markup=reply_markup
    )
    
    # Foydalanuvchi ma'lumotlarini saqlash uchun
    context.user_data['user_info'] = {
        'telegram_username': user.username or 'Username yo\'q',
        'telegram_id': user.id
    }
    
    return FISH

async def start_form_callback(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Admin uchun ariza to'ldirish boshlash"""
    query = update.callback_query
    await query.answer()
    
    user = query.from_user
    
    await query.message.reply_text(
        f"Assalomu alaykum, {user.first_name}!\n\n"
        "🧪 Kimyo kursimizga yozilishingiz uchun so'rovnomalarga "
        "aniq va to'la javob berishingiz kerak.\n\n"
        "📝 Iltimos, to'liq ismingizni kiriting (F.I.SH):",
        reply_markup=ReplyKeyboardRemove()
    )
    
    context.user_data['user_info'] = {
        'telegram_username': user.username or 'Username yo\'q',
        'telegram_id': user.id
    }
    
    return FISH

async def show_admin_panel_callback(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Callback orqali admin panelni ko'rsatish"""
    query = update.callback_query
    await query.answer()
    
    users_count = get_users_count()
    unread_count = len(get_unread_messages())
    keyboard = [
        [InlineKeyboardButton("📊 Statistika", callback_data='admin_stats')],
        [InlineKeyboardButton("📢 Xabar yuborish", callback_data='admin_broadcast')],
        [InlineKeyboardButton("👥 Foydalanuvchilar", callback_data='admin_users_list')],
        [InlineKeyboardButton("📝 Davomat", callback_data='admin_attendance')],
        [InlineKeyboardButton(f"💬 Murojaatlar ({unread_count})", callback_data='admin_messages')],
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await query.message.reply_text(
        f"🔐 <b>ADMIN PANEL</b>\n\n"
        f"👥 Jami foydalanuvchilar: <b>{users_count}</b>\n"
        f"💬 Yangi murojaatlar: <b>{unread_count}</b>\n"
        f"📅 Bugun: {datetime.now().strftime('%d.%m.%Y')}\n\n"
        f"Quyidagi tugmalardan birini tanlang:",
        reply_markup=reply_markup,
        parse_mode='HTML'
    )

async def get_fish(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """F.I.SH ni qabul qilish"""
    context.user_data['user_info']['fish'] = update.message.text
    await update.message.reply_text("📍 Qaysi tuman, qaysi qishloqdansiz?")
    return MANZIL

async def get_manzil(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Manzilni qabul qilish"""
    context.user_data['user_info']['manzil'] = update.message.text
    await update.message.reply_text("🎓 Nechanchi sinfsiz yoki bitirganmisiz?")
    return SINF

async def get_sinf(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Sinf ma'lumotini qabul qilish"""
    context.user_data['user_info']['sinf'] = update.message.text
    
    keyboard = [['Ha', 'Yo\'q']]
    reply_markup = ReplyKeyboardMarkup(keyboard, one_time_keyboard=True, resize_keyboard=True)
    
    await update.message.reply_text(
        "📚 Avval kimyo o'qiganmisiz?",
        reply_markup=reply_markup
    )
    return AVVAL_OQIGAN

async def get_avval_oqigan(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Avval o'qigan ma'lumotini qabul qilish"""
    context.user_data['user_info']['avval_oqigan'] = update.message.text
    
    keyboard = [['Ha', 'Yo\'q']]
    reply_markup = ReplyKeyboardMarkup(keyboard, one_time_keyboard=True, resize_keyboard=True)
    
    await update.message.reply_text(
        "👨‍👩‍👦 Ota-onangiz bormi?",
        reply_markup=reply_markup
    )
    return OTA_ONA

async def get_ota_ona(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Ota-ona ma'lumotini qabul qilish"""
    context.user_data['user_info']['ota_ona'] = update.message.text
    
    await update.message.reply_text(
        "📱 O'zingizning telefon raqamingizni kiriting:\n"
        "(Masalan: +998901234567)",
        reply_markup=ReplyKeyboardRemove()
    )
    return TELEFON

async def get_telefon(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Telefon raqamni qabul qilish"""
    context.user_data['user_info']['telefon'] = update.message.text
    
    await update.message.reply_text(
        "📞 Ota-onangizning telefon raqamini kiriting:\n"
        "(Masalan: +998901234567)"
    )
    return OTA_ONA_TELEFON

async def get_ota_ona_telefon(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Ota-ona telefon raqamini qabul qilish"""
    context.user_data['user_info']['ota_ona_telefon'] = update.message.text
    
    await update.message.reply_text(
        "📅 Kimyo kursiga qachon kelgansiz?\n"
        "(Masalan: 15-yanvar 2026 yoki Endi bormoqchiman)"
    )
    return KELGAN_SANA

async def get_kelgan_sana(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Kelgan sanani qabul qilish"""
    context.user_data['user_info']['kelgan_sana'] = update.message.text
    
    await update.message.reply_text(
        "📸 Rasmingizni yuboring.\n\n"
        "⚠️ Shaxsiy ma'lumotlaringiz sir saqlanadi!!!"
    )
    return RASM

async def get_rasm(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Rasmni qabul qilish"""
    if update.message.photo:
        photo = update.message.photo[-1]
        
        # Rasmni saqlash
        context.user_data['user_info']['photo_file_id'] = photo.file_id
        
        logger.info(f"Rasm qabul qilindi: {photo.file_id}")
        
        await update.message.reply_text(
            "✅ Rasm qabul qilindi!\n\n"
            "🎯 Kimyo o'qishdan maqsadingiz nima?"
        )
        return MAQSAD
    else:
        await update.message.reply_text(
            "❌ Iltimos, rasm yuboring!\n\n"
            "📸 Rasmingizni qayta yuboring:"
        )
        return RASM

async def get_maqsad(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Maqsadni qabul qilish va ma'lumotlarni adminga yuborish"""
    context.user_data['user_info']['maqsad'] = update.message.text
    user_info = context.user_data['user_info']
    
    logger.info(f"Ma'lumotlar to'plandi: {user_info.get('fish')}")
    
    # Foydalanuvchini bazaga qo'shish
    add_user(
        user_info.get('telegram_id'),
        user_info.get('telegram_username'),
        user_info.get('fish')
    )
    
    try:
        # Word hujjat yaratish
        doc = Document()
        
        # Sarlavha
        title = doc.add_heading('KIMYO KURSI - YANGI TALABA MA\'LUMOTLARI', 0)
        title.alignment = 1  # Center
        
        doc.add_paragraph()
        
        # Ma'lumotlarni qo'shish
        data_pairs = [
            ("👤 F.I.SH:", user_info.get('fish', 'N/A')),
            ("📍 Manzil:", user_info.get('manzil', 'N/A')),
            ("🎓 Sinf:", user_info.get('sinf', 'N/A')),
            ("📚 Avval kimyo o'qiganmi:", user_info.get('avval_oqigan', 'N/A')),
            ("👨‍👩‍👦 Ota-onasi bormi:", user_info.get('ota_ona', 'N/A')),
            ("📱 Telefon raqami:", user_info.get('telefon', 'N/A')),
            ("📞 Ota-ona telefon raqami:", user_info.get('ota_ona_telefon', 'N/A')),
            ("📅 Kelgan sana:", user_info.get('kelgan_sana', 'N/A')),
            ("🎯 Maqsadi:", user_info.get('maqsad', 'N/A')),
        ]
        
        for label, value in data_pairs:
            p = doc.add_paragraph()
            p.add_run(label).bold = True
            p.add_run(f" {value}")
        
        # Rasm haqida ma'lumot
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("📸 Rasm: ").bold = True
        p.add_run("Telegram orqali yuborilgan (quyida)")
        
        doc.add_paragraph()
        doc.add_paragraph("_" * 50)
        
        # Telegram ma'lumotlari
        p = doc.add_paragraph()
        p.add_run("📱 Telegram Username: ").bold = True
        p.add_run(f"@{user_info.get('telegram_username', 'N/A')}")
        
        p = doc.add_paragraph()
        p.add_run("🆔 Telegram ID: ").bold = True
        p.add_run(str(user_info.get('telegram_id', 'N/A')))
        
        # Hujjatni xotiraga saqlash
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        # Faylni adminga yuborish
        file_name = f"Ariza_{user_info.get('fish', 'Nomsiz').replace(' ', '_')}.docx"
        logger.info("Adminga fayl yuborilmoqda...")
        
        # Word faylni yuborish
        await context.bot.send_document(
            chat_id=ADMIN_ID,
            document=file_stream,
            filename=file_name,
            caption=f"🆕 Yangi ariza keldi!\n\n"
                    f"👤 {user_info.get('fish', 'N/A')}\n"
                    f"📱 @{user_info.get('telegram_username', 'N/A')}\n"
                    f"📅 Kelgan sana: {user_info.get('kelgan_sana', 'N/A')}"
        )
        
        logger.info("Word fayl yuborildi")
        
        # Rasmni yuborish
        if 'photo_file_id' in user_info:
            logger.info("Rasm yuborilmoqda...")
            await context.bot.send_photo(
                chat_id=ADMIN_ID,
                photo=user_info.get('photo_file_id'),
                caption=f"📸 {user_info.get('fish', 'N/A')} - Rasm"
            )
            logger.info("Rasm yuborildi")
        
        await update.message.reply_text(
            "✅ Arizangiz muvaffaqiyatli qabul qilindi!\n\n"
            "Ma'lumotlaringiz admin ko'rib chiqish uchun yuborildi. "
            "Tez orada siz bilan bog'lanamiz.\n\n"
            "Rahmat! 🙏"
        )
        
        logger.info("Jarayon muvaffaqiyatli yakunlandi")
        
    except Exception as e:
        logger.error(f"Xatolik yuz berdi: {e}", exc_info=True)
        await update.message.reply_text(
            "❌ Arizani yuborishda xatolik yuz berdi. "
            "Iltimos, qaytadan urinib ko'ring yoki admin bilan bog'laning."
        )
    
    return ConversationHandler.END

async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """So'rovnomani bekor qilish"""
    await update.message.reply_text(
        "❌ So'rovnoma bekor qilindi.\n"
        "Qayta boshlash uchun /start ni bosing.",
        reply_markup=ReplyKeyboardRemove()
    )
    return ConversationHandler.END

async def error_handler(update: object, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Error handlerlarni log qilish"""
    logger.error(f"Xatolik yuz berdi: {context.error}", exc_info=context.error)

def main() -> None:
    """Botni ishga tushirish"""
    try:
        # Database yaratish
        logger.info("Database yaratilmoqda...")
        init_db()
        logger.info("Database tayyor!")
        
        application = Application.builder().token(BOT_TOKEN).build()
        
        # Error handler qo'shish
        application.add_error_handler(error_handler)
        
        # Admin panel handler
        application.add_handler(CommandHandler('admin', admin_panel))
        
        # Admin callback handler
        application.add_handler(CallbackQueryHandler(show_admin_panel_callback, pattern='show_admin_panel'))
        application.add_handler(CallbackQueryHandler(start_form_callback, pattern='start_form'))
        application.add_handler(CallbackQueryHandler(check_in_callback, pattern='check_in'))
        
        # Broadcast conversation handler - MATN, RASM, VIDEO, DOCUMENT qabul qiladi
        broadcast_handler = ConversationHandler(
            entry_points=[CallbackQueryHandler(admin_callback, pattern='admin_broadcast')],
            states={
                ADMIN_BROADCAST: [
                    MessageHandler(filters.TEXT & ~filters.COMMAND, broadcast_message),
                    MessageHandler(filters.PHOTO, broadcast_message),
                    MessageHandler(filters.VIDEO, broadcast_message),
                    MessageHandler(filters.Document.ALL, broadcast_message),
                ],
            },
            fallbacks=[CommandHandler('cancel', cancel_broadcast)],
        )
        application.add_handler(broadcast_handler)
        
        # Contact admin handler
        contact_handler = ConversationHandler(
            entry_points=[CallbackQueryHandler(contact_admin_start, pattern='contact_admin')],
            states={
                ADMIN_MESSAGE: [MessageHandler(filters.TEXT & ~filters.COMMAND, receive_admin_message)],
            },
            fallbacks=[CommandHandler('cancel', cancel)],
        )
        application.add_handler(contact_handler)
        
        # Admin callbacks
        application.add_handler(CallbackQueryHandler(admin_callback, pattern='admin_stats|admin_users_list|admin_attendance|admin_messages|start_attendance|end_attendance|back_to_admin'))
        
        # Main conversation handler - Foydalanuvchilar uchun
        conv_handler = ConversationHandler(
            entry_points=[CommandHandler('start', start)],
            states={
                FISH: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_fish)],
                MANZIL: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_manzil)],
                SINF: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_sinf)],
                AVVAL_OQIGAN: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_avval_oqigan)],
                OTA_ONA: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_ota_ona)],
                TELEFON: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_telefon)],
                OTA_ONA_TELEFON: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_ota_ona_telefon)],
                KELGAN_SANA: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_kelgan_sana)],
                RASM: [MessageHandler(filters.PHOTO, get_rasm)],
                MAQSAD: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_maqsad)],
            },
            fallbacks=[CommandHandler('cancel', cancel)],
        )
        
        application.add_handler(conv_handler)
        
        # Botni ishga tushirish
        logger.info("Bot ishga tushirilmoqda...")
        logger.info(f"Database path: {DB_PATH}")
        application.run_polling(allowed_updates=Update.ALL_TYPES, drop_pending_updates=True)
        
    except Exception as e:
        logger.error(f"Botni ishga tushirishda xatolik: {e}", exc_info=True)
        raise

if __name__ == '__main__':
    main()
